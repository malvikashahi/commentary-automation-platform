import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.models.schemas import GeneratedCommentary, ExportResult


EXPORT_DIR = "outputs"
os.makedirs(EXPORT_DIR, exist_ok=True)


class ExportEngine:
    """Exports generated commentary to DOCX and PDF formats."""

    def export_docx(
        self, commentary: GeneratedCommentary
    ) -> ExportResult:
        doc = Document()

        # ── Styles ──────────────────────────────────────────
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10.5)

        # ── Cover Page ──────────────────────────────────────
        self._add_cover(doc, commentary)

        # ── Sections ────────────────────────────────────────
        for section in commentary.sections:
            # Section heading
            heading = doc.add_heading(section.section_name, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0x0A, 0x22, 0x40)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].font.bold = True

            # Section content
            for para_text in section.content.split('\n'):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.space_before = Pt(0)

            doc.add_paragraph()  # Spacer

        # ── Footer ──────────────────────────────────────────
        self._add_footer(doc, commentary)

        # ── Save ────────────────────────────────────────────
        export_id = str(uuid.uuid4())[:8]
        filename = (
            f"{commentary.account_id}_{commentary.period}_"
            f"Commentary_{export_id}.docx"
        )
        file_path = os.path.join(EXPORT_DIR, filename)
        doc.save(file_path)

        return ExportResult(
            export_id=export_id,
            file_path=file_path,
            format="docx",
            account_count=1
        )

    def _add_cover(self, doc: Document, commentary: GeneratedCommentary):
        doc.add_heading('Portfolio Commentary', 0).runs[0].font.color.rgb = (
            RGBColor(0x0A, 0x22, 0x40)
        )
        doc.add_paragraph(f"Account: {commentary.account_id}")
        doc.add_paragraph(f"Strategy: {commentary.strategy_id}")
        doc.add_paragraph(f"Period: {commentary.period}")
        doc.add_paragraph(
            f"Generated: {commentary.generated_at.strftime('%B %d, %Y')}"
        )
        doc.add_paragraph(
            "FOR INSTITUTIONAL USE ONLY — NOT FOR PUBLIC DISTRIBUTION"
        ).runs[0].italic = True
        doc.add_page_break()

    def _add_footer(self, doc: Document, commentary: GeneratedCommentary):
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = (
            f"T. Rowe Price | {commentary.strategy_id} | "
            f"{commentary.period} | Confidential"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x9A, 0xA3, 0xAF)