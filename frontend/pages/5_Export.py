# frontend/pages/5_Export.py

import streamlit as st
import requests
import os
import time
from datetime import datetime

API_BASE   = "http://localhost:8000/api"
OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Custom CSS ───────────────────────────────────────────────
st.markdown("""
<style>
.export-card {
    background: white;
    border: 1px solid #DDE1E7;
    border-radius: 10px;
    padding: 20px 24px;
    margin-bottom: 12px;
    box-shadow: 0 2px 8px rgba(10,34,64,0.06);
}
.format-badge-docx {
    background: #E8F4FD; color: #1A6FB0;
    padding: 3px 10px; border-radius: 99px;
    font-size: 0.75rem; font-weight: 700;
}
.format-badge-pdf {
    background: #FEF0EF; color: #D93025;
    padding: 3px 10px; border-radius: 99px;
    font-size: 0.75rem; font-weight: 700;
}
</style>
""", unsafe_allow_html=True)

# ── Page Header ──────────────────────────────────────────────
st.markdown("""
<div style="background:linear-gradient(135deg,#004B49,#0c3635);
            padding:24px 28px;border-radius:10px;
            margin-bottom:20px;border-left:5px solid #00A99D;">
    <div style="color:white;font-size:1.5rem;
                font-weight:800;margin:0;">
        📤 Export Center
    </div>
    <div style="color:rgba(255,255,255,0.6);
                margin:6px 0 0 0;font-size:0.85rem;">
        Download generated commentaries as DOCX or PDF
    </div>
</div>
""", unsafe_allow_html=True)

# ── Check reportlab availability upfront ────────────────────
try:
    import reportlab
    REPORTLAB_AVAILABLE = False
try:
    import reportlab
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

if not REPORTLAB_AVAILABLE:
    st.warning(
        "⚠️ **reportlab not installed** — "
        "PDF export will use DOCX fallback. "
        "To enable PDF: run `pip install reportlab` "
        "in your venv terminal, then restart Streamlit."
    )

# ── Check if commentary exists ───────────────────────────────
has_commentary = "generated_commentary" in st.session_state
has_approved   = (
    has_commentary and
    st.session_state["generated_commentary"]
    .get("commentary") and
    st.session_state["generated_commentary"]
    ["commentary"].get("status") == "APPROVED"
)

account_id  = ""
strategy_id = ""
period      = ""
sections    = []

if has_commentary:
    _comm       = (
        st.session_state["generated_commentary"]
        .get("commentary") or {}
    )
    account_id  = _comm.get("account_id")  or ""
    strategy_id = _comm.get("strategy_id") or ""
    period      = _comm.get("period")      or "4Q25"
    sections    = _comm.get("sections")    or []

# ── Status Overview ──────────────────────────────────────────
st.subheader("📊 Export Status")

c1, c2, c3, c4 = st.columns(4)
with c1:
    st.metric(
        "Commentary Ready",
        "✅ Yes" if has_commentary else "❌ No"
    )
with c2:
    st.metric(
        "Approval Status",
        "✅ Approved" if has_approved else "⏳ Pending"
    )
with c3:
    st.metric("Account", account_id or "—")
with c4:
    st.metric("Period", period or "—")

st.markdown("---")

# ── Export Options ───────────────────────────────────────────
st.subheader("📎 Export Options")

if not has_commentary:
    st.warning(
        "⚠️ No commentary available. "
        "Please complete "
        "**Upload → Validation → Processing → Review** first."
    )
    st.stop()

if not has_approved:
    st.warning(
        "⚠️ Commentary not yet approved. "
        "Please go to **Review** page "
        "and click **Approve Commentary** first."
    )
    st.info("💡 You can still export a **draft** version below.")

# ── Format Cards ─────────────────────────────────────────────
col_fmt1, col_fmt2 = st.columns(2)

with col_fmt1:
    st.markdown("""
    <div class="export-card">
        <div style="font-size:2rem;margin-bottom:8px;">📝</div>
        <div style="font-weight:700;color:#004B49;font-size:1rem;">
            Word Document
        </div>
        <span class="format-badge-docx">.DOCX</span>
        <div style="font-size:0.8rem;color:#9AA3AF;margin-top:8px;">
            Editable Microsoft Word format.<br/>
            Best for further editing and review.
        </div>
    </div>
    """, unsafe_allow_html=True)
    export_docx = st.button(
        "📥 Export as DOCX",
        type="primary",
        use_container_width=True,
        key="btn_docx"
    )

with col_fmt2:
    pdf_label = (
        "📥 Export as PDF"
        if REPORTLAB_AVAILABLE
        else "📥 Export as PDF (DOCX fallback)"
    )
    st.markdown(f"""
    <div class="export-card">
        <div style="font-size:2rem;margin-bottom:8px;">📄</div>
        <div style="font-weight:700;color:#004B49;font-size:1rem;">
            PDF Document
        </div>
        <span class="format-badge-pdf">
            {"PDF" if REPORTLAB_AVAILABLE else "PDF → DOCX fallback"}
        </span>
        <div style="font-size:0.8rem;color:#9AA3AF;margin-top:8px;">
            {"Fixed-format PDF for distribution." 
             if REPORTLAB_AVAILABLE 
             else "Install reportlab for true PDF export."}
        </div>
    </div>
    """, unsafe_allow_html=True)
    export_pdf = st.button(
        pdf_label,
        use_container_width=True,
        key="btn_pdf"
    )

st.markdown("---")

# ════════════════════════════════════════════════════════════
# DOCX EXPORT
# ════════════════════════════════════════════════════════════
if export_docx:
    with st.spinner("📝 Generating Word document..."):
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = DocxDocument()

            # Cover page
            title_para = doc.add_heading(
                'Portfolio Commentary', 0
            )
            if title_para.runs:
                title_para.runs[0].font.color.rgb = RGBColor(
                    0x00, 0x4B, 0x49
                )
                title_para.runs[0].font.size = Pt(24)

            doc.add_paragraph(f"Account:   {account_id}")
            doc.add_paragraph(f"Strategy:  {strategy_id}")
            doc.add_paragraph(f"Period:    {period}")
            doc.add_paragraph(
                f"Generated: "
                f"{datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
            )
            doc.add_paragraph(
                f"Status:    "
                f"{'APPROVED' if has_approved else 'DRAFT'}"
            )
            disc = doc.add_paragraph(
                "FOR INSTITUTIONAL USE ONLY — "
                "NOT FOR PUBLIC DISTRIBUTION"
            )
            if disc.runs:
                disc.runs[0].italic = True
                disc.runs[0].font.color.rgb = RGBColor(
                    0x9A, 0xA3, 0xAF
                )
            doc.add_page_break()

            # Commentary sections
            for section in sections:
                sname   = section.get("section_name") or "Section"
                content = section.get("content")      or ""
                is_auto = section.get("is_auto_generated", True)

                heading = doc.add_heading(sname, level=2)
                if heading.runs:
                    heading.runs[0].font.color.rgb = RGBColor(
                        0x00, 0x4B, 0x49
                    )
                    heading.runs[0].font.size = Pt(12)
                    heading.runs[0].bold = True

                if is_auto:
                    lbl = doc.add_paragraph("[ Auto-Generated ]")
                    if lbl.runs:
                        lbl.runs[0].italic = True
                        lbl.runs[0].font.size = Pt(8)
                        lbl.runs[0].font.color.rgb = RGBColor(
                            0x9A, 0xA3, 0xAF
                        )

                for line in content.split('\n'):
                    if line.strip():
                        p = doc.add_paragraph(line.strip())
                        p.paragraph_format.space_after  = Pt(4)
                        p.paragraph_format.space_before = Pt(0)
                        if p.runs:
                            p.runs[0].font.size = Pt(10.5)

                doc.add_paragraph()

            # Footer
            sec_obj = doc.sections[0]
            footer  = sec_obj.footer
            fp      = footer.paragraphs[0]
            fp.text = (
                f"Acuity Analytics | {strategy_id} | "
                f"{period} | Confidential | "
                f"Generated {datetime.now().strftime('%d %b %Y')}"
            )
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if fp.runs:
                fp.runs[0].font.size = Pt(8)
                fp.runs[0].font.color.rgb = RGBColor(
                    0x9A, 0xA3, 0xAF
                )

            # Save
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename  = (
                f"{account_id}_{strategy_id}_{period}_"
                f"Commentary_{timestamp}.docx"
            )
            file_path = os.path.join(OUTPUT_DIR, filename)
            doc.save(file_path)

            time.sleep(0.3)
            st.success(f"✅ DOCX generated: **{filename}**")

            with open(file_path, "rb") as f:
                docx_bytes = f.read()

            st.download_button(
                label="📥 Click Here to Download DOCX",
                data=docx_bytes,
                file_name=filename,
                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                ),
                use_container_width=True
            )

            if "export_history" not in st.session_state:
                st.session_state["export_history"] = []
            st.session_state["export_history"].append({
                "filename":  filename,
                "format":    "DOCX",
                "account":   account_id,
                "period":    period,
                "timestamp": datetime.now().strftime(
                    "%Y-%m-%d %H:%M:%S"
                ),
                "path": file_path
            })

        except Exception as e:
            st.error(f"❌ DOCX Export failed: {str(e)}")
            st.exception(e)

# ════════════════════════════════════════════════════════════
# PDF EXPORT
# ════════════════════════════════════════════════════════════
if export_pdf:
    with st.spinner("📄 Generating PDF document..."):
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            if REPORTLAB_AVAILABLE:
                # ── True PDF via reportlab ────────────────────
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import (
                    getSampleStyleSheet, ParagraphStyle
                )
                from reportlab.lib.units import inch
                from reportlab.lib import colors
                from reportlab.platypus import (
                    SimpleDocTemplate, Paragraph,
                    Spacer, HRFlowable, PageBreak
                )

                filename  = (
                    f"{account_id}_{strategy_id}_{period}_"
                    f"Commentary_{timestamp}.pdf"
                )
                file_path = os.path.join(OUTPUT_DIR, filename)

                navy = colors.HexColor('#004B49')
                teal = colors.HexColor('#00A99D')
                gray = colors.HexColor('#9AA3AF')
                dark = colors.HexColor('#2C3542')

                styles = getSampleStyleSheet()

                title_style = ParagraphStyle(
                    'PTitle',
                    parent=styles['Title'],
                    fontSize=22,
                    textColor=navy,
                    spaceAfter=12,
                    fontName='Helvetica-Bold'
                )
                heading_style = ParagraphStyle(
                    'PHeading',
                    parent=styles['Heading2'],
                    fontSize=12,
                    textColor=navy,
                    spaceBefore=14,
                    spaceAfter=6,
                    fontName='Helvetica-Bold'
                )
                body_style = ParagraphStyle(
                    'PBody',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=dark,
                    spaceAfter=6,
                    leading=16,
                    fontName='Helvetica'
                )
                meta_style = ParagraphStyle(
                    'PMeta',
                    parent=styles['Normal'],
                    fontSize=9,
                    textColor=gray,
                    spaceAfter=4,
                    fontName='Helvetica'
                )
                disc_style = ParagraphStyle(
                    'PDisc',
                    parent=styles['Normal'],
                    fontSize=8,
                    textColor=gray,
                    fontName='Helvetica-Oblique'
                )

                doc_pdf = SimpleDocTemplate(
                    file_path,
                    pagesize=A4,
                    rightMargin=inch,
                    leftMargin=inch,
                    topMargin=inch,
                    bottomMargin=inch
                )

                elems = []
                elems.append(
                    Paragraph("Portfolio Commentary", title_style)
                )
                elems.append(
                    HRFlowable(
                        width="100%", thickness=2,
                        color=teal, spaceAfter=12
                    )
                )
                elems.append(
                    Paragraph(f"Account: {account_id}", meta_style)
                )
                elems.append(
                    Paragraph(
                        f"Strategy: {strategy_id}", meta_style
                    )
                )
                elems.append(
                    Paragraph(f"Period: {period}", meta_style)
                )
                elems.append(
                    Paragraph(
                        f"Generated: "
                        f"{datetime.now().strftime('%B %d, %Y')}",
                        meta_style
                    )
                )
                elems.append(
                    Paragraph(
                        f"Status: "
                        f"{'APPROVED' if has_approved else 'DRAFT'}",
                        meta_style
                    )
                )
                elems.append(Spacer(1, 10))
                elems.append(
                    Paragraph(
                        "FOR INSTITUTIONAL USE ONLY — "
                        "NOT FOR PUBLIC DISTRIBUTION",
                        disc_style
                    )
                )
                elems.append(PageBreak())

                for section in sections:
                    sname   = (
                        section.get("section_name") or "Section"
                    )
                    content = section.get("content") or ""

                    elems.append(
                        Paragraph(sname.upper(), heading_style)
                    )
                    elems.append(
                        HRFlowable(
                            width="100%", thickness=1,
                            color=teal, spaceAfter=8
                        )
                    )
                    for line in content.split('\n'):
                        if line.strip():
                            safe = (
                                line.strip()
                                .replace('&', '&amp;')
                                .replace('<', '&lt;')
                                .replace('>', '&gt;')
                            )
                            elems.append(
                                Paragraph(safe, body_style)
                            )
                    elems.append(Spacer(1, 12))

                doc_pdf.build(elems)
                mime     = "application/pdf"
                dl_label = "📥 Click Here to Download PDF"
                fmt_name = "PDF"

            else:
                # ── DOCX fallback when reportlab missing ──────
                from docx import Document as DocxDocument
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                filename  = (
                    f"{account_id}_{strategy_id}_{period}_"
                    f"Commentary_{timestamp}_PDF_FALLBACK.docx"
                )
                file_path = os.path.join(OUTPUT_DIR, filename)

                doc = DocxDocument()
                doc.add_heading('Portfolio Commentary', 0)
                doc.add_paragraph(f"Account:   {account_id}")
                doc.add_paragraph(f"Strategy:  {strategy_id}")
                doc.add_paragraph(f"Period:    {period}")
                doc.add_paragraph(
                    f"Generated: "
                    f"{datetime.now().strftime('%B %d, %Y')}"
                )
                doc.add_paragraph(
                    "NOTE: Install reportlab for true PDF export."
                )
                doc.add_page_break()

                for section in sections:
                    sname   = (
                        section.get("section_name") or "Section"
                    )
                    content = section.get("content") or ""
                    doc.add_heading(sname, level=2)
                    for line in content.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                    doc.add_paragraph()

                doc.save(file_path)
                mime     = (
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                )
                dl_label = (
                    "📥 Download DOCX "
                    "(install reportlab for true PDF)"
                )
                fmt_name = "DOCX"

            # ── Download button ───────────────────────────────
            time.sleep(0.3)
            st.success(f"✅ File generated: **{filename}**")

            with open(file_path, "rb") as f:
                file_bytes = f.read()

            st.download_button(
                label=dl_label,
                data=file_bytes,
                file_name=filename,
                mime=mime,
                use_container_width=True
            )

            if "export_history" not in st.session_state:
                st.session_state["export_history"] = []
            st.session_state["export_history"].append({
                "filename":  filename,
                "format":    fmt_name,
                "account":   account_id,
                "period":    period,
                "timestamp": datetime.now().strftime(
                    "%Y-%m-%d %H:%M:%S"
                ),
                "path": file_path
            })

        except Exception as e:
            st.error(f"❌ PDF Export failed: {str(e)}")
            st.exception(e)

st.markdown("---")

# ════════════════════════════════════════════════════════════
# EXPORT HISTORY
# ════════════════════════════════════════════════════════════
st.subheader("📋 Export History")

history = st.session_state.get("export_history", [])

if not history:
    st.info(
        "No exports yet. "
        "Generate a DOCX or PDF above to see history here."
    )
else:
    for i, item in enumerate(reversed(history)):
        fmt   = item.get("format", "DOCX")
        badge = (
            "<span class='format-badge-docx'>DOCX</span>"
            if fmt == "DOCX"
            else "<span class='format-badge-pdf'>PDF</span>"
        )
        st.markdown(f"""
        <div class="export-card">
            <div style="display:flex;justify-content:space-between;
                        align-items:center;">
                <div>
                    {badge}
                    <span style="font-weight:700;color:#004B49;
                                 margin-left:8px;font-size:0.9rem;">
                        {item.get("filename", "Unknown")}
                    </span>
                </div>
                <div style="font-size:0.78rem;color:#9AA3AF;">
                    {item.get("timestamp", "")}
                </div>
            </div>
            <div style="font-size:0.78rem;color:#9AA3AF;
                        margin-top:6px;">
                Account: <b>{item.get("account", "")}</b> |
                Period: <b>{item.get("period", "")}</b>
            </div>
        </div>
        """, unsafe_allow_html=True)

        fp = item.get("path", "")
        if fp and os.path.exists(fp):
            with open(fp, "rb") as f:
                fb = f.read()
            mime_type = (
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
                if fmt == "DOCX"
                else "application/pdf"
            )
            st.download_button(
                label=f"📥 Re-download {fmt}",
                data=fb,
                file_name=item.get(
                    "filename", f"export.{fmt.lower()}"
                ),
                mime=mime_type,
                key=f"redl_{i}_{item.get('timestamp', i)}"
            )

st.markdown("---")

# ── Footer ────────────────────────────────────────────────────
st.markdown(
    "<div style='text-align:center;padding:8px 0;'>"
    "<span style='color:#9AA3AF;font-size:0.75rem;'>"
    "Acuity Analytics &middot; "
    "Commentary Automation Platform v1.0 &middot; Confidential"
    "</span></div>",
    unsafe_allow_html=True
)